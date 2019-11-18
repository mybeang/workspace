import pdb
import re
import os
import sys
import json
import yaml
import optparse
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from pandas import DataFrame
from string import Template
from collections import OrderedDict


VERSION = "0.0.1"


class Docx(object):
    def __init__(self, detail_log_dir):
        with open(detail_log_dir + "summary_result.json") as f:
            self.result_data = json.load(f, object_pairs_hook=OrderedDict)

        self.detail_log_list =  os.listdir(detail_log_dir)
        self.detail_log_list.remove("summary_result.json")

        current_path = os.getcwd()
        if sys.platform == 'win32' or sys.platform == 'cygwin':
            template = "\\Evt\\template.docx"
        else:
            template = "/Evt/template.docx"
        self.doc = Document(current_path+template)
        self.report_string = ""

    def make_summary_result(self):
        self.doc.add_heading("Summary Result", level=1)
        indexs = self.result_data.keys()
        values = [self.result_data[k] for k in indexs]
        df = DataFrame(values, index=indexs)
        t = self.doc.add_table(df.shape[0]+1, df.shape[1]+1)
        t.style = "summary_result2"

        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        # add the header rows.
        t.cell(0, 0).text = "TestCase"
        for column_index in range(df.shape[-1]):
            text = str(df.columns[column_index]+1)
            t.cell(0, column_index + 1).text = text
            par = t.cell(0, column_index + 1).paragraphs[0]
            par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # add the rest of the data frame
        for row_index in range(df.shape[0]):
            t.cell(row_index + 1, 0).text = df.index[row_index]
            for column_index in range(df.shape[-1]):
                text = str(df.values[row_index, column_index])
                if text.lower() == "none":
                    text = ""
                t.cell(row_index + 1, column_index + 1).text = text
                par = t.cell(row_index + 1, column_index + 1).paragraphs[0]
                par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                font = par.runs[0].font
                if text.lower() == 'pass':
                    font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
                elif text.lower() == 'fail':
                    font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


    def prepare_docx(self, data):
        device_model = data["Device"]["Name"]
        device_description = data["Device"]["Description"]
        front_page = {
            "device_model": device_model,
            "device_description": device_description
        }
        for i, paragraph in enumerate(self.doc.paragraphs):
            text = Template(paragraph.text).substitute(**front_page)
            paragraph.text = text

    def add_text(self, string_data_list):
        for string_data in string_data_list:
            try:
                date, time, level, msg = string_data.split(" ", 3)
                time = re.sub(r',\d+', '', time)
                msg = re.sub(r'{.+}', '', msg)
                msg = msg.strip()
            except:
                date, time, level, msg = [""] * 4

            if re.findall(r'\d+-\d+-\d+', date) and re.findall(r"\d+:\d+:\d+", time):
                #self._print_report(self.report_string)
                self.report_string = ""
                #text = "{} {} {} {}".format(date, time, level, msg)
                #self._sel_string(level, msg)
            else:
                self.report_string += string_data

    def make_file(self, path, device):
        self.doc.save(path + "{}_test_result.docx".format(device))

def main():
    if sys.platform == 'win32' or sys.platform == 'cygwin':
        home_path = os.path.expanduser("~\Documents\\")
    else:
        home_path = "~/"

    usage = "Usage: %prog DetailLogPath"
    version = "%prog " + VERSION
    parser = optparse.OptionParser(usage=usage, version=version)
    options, args = parser.parse_args()

    if len(args) == 0:
        print("No Path")
        return
    elif len(args) > 1:
        print("Not support multi path")
        return

    if not os.path.exists(args[0]):
        print("No exist the path: {}".format(args[0]))
        return

    docx = Docx(args[0])
    with open(home_path + "evt_conf.yaml", "r") as f:
        test_conf = yaml.load(f)
    docx.prepare_docx(test_conf)
    docx.make_summary_result()
    docx.make_file(args[0], test_conf['Device']['Name'])
